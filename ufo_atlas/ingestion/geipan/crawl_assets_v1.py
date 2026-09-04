#!/usr/bin/env python3
"""Preserve and extract all unique public assets linked from GEIPAN case pages.

Raw assets are retained in per-batch source shards. The aggregate database keeps
complete retrieval metadata, hashes, source links, and extracted text/structure
without combining all binaries into one oversized Atlas import package.
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import io
import json
import mimetypes
import shutil
import sqlite3
import subprocess
import time
import urllib.parse
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

import crawl_case_packets as shared


def clean(value: Any) -> str:
    return shared.clean(value)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def extension_for(url: str, content_type: str | None) -> str:
    ext = Path(urllib.parse.urlparse(url).path).suffix.lower()
    if ext and len(ext) <= 12:
        return ext
    mime = (content_type or "").split(";", 1)[0].strip().lower()
    return mimetypes.guess_extension(mime) or ".bin"


def decode_text(data: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), "utf-8-replace"


def extract_pdf(data: bytes) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    from pypdf import PdfReader

    pages: list[dict[str, Any]] = []
    metadata: dict[str, Any] = {}
    try:
        reader = PdfReader(io.BytesIO(data), strict=False)
        metadata["page_count"] = len(reader.pages)
        metadata["document_metadata"] = {str(k): clean(v) for k, v in (reader.metadata or {}).items()}
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
                error = None
            except Exception as exc:  # noqa: BLE001
                text = ""
                error = f"{type(exc).__name__}: {exc}"
            pages.append(
                {
                    "unit_number": page_number,
                    "unit_type": "PDF_PAGE",
                    "text": text,
                    "text_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                    "error": error,
                }
            )
        metadata["text_status"] = "TEXT_LAYER_PRESENT" if any(clean(p["text"]) for p in pages) else "NO_TEXT_LAYER"
    except Exception as exc:  # noqa: BLE001
        metadata["parse_error"] = f"{type(exc).__name__}: {exc}"
        metadata["text_status"] = "PDF_PARSE_ERROR"
    return pages, metadata


def extract_image(data: bytes) -> dict[str, Any]:
    from PIL import Image, ExifTags

    result: dict[str, Any] = {}
    try:
        with Image.open(io.BytesIO(data)) as image:
            result.update(
                {
                    "format": image.format,
                    "width": image.width,
                    "height": image.height,
                    "mode": image.mode,
                    "frame_count": getattr(image, "n_frames", 1),
                }
            )
            try:
                exif = image.getexif()
                result["exif"] = {
                    str(ExifTags.TAGS.get(key, key)): clean(value)
                    for key, value in exif.items()
                    if clean(value)
                }
            except Exception as exc:  # noqa: BLE001
                result["exif_error"] = f"{type(exc).__name__}: {exc}"
    except Exception as exc:  # noqa: BLE001
        result["parse_error"] = f"{type(exc).__name__}: {exc}"
    return result


def extract_docx(data: bytes) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    from docx import Document

    units: list[dict[str, Any]] = []
    result: dict[str, Any] = {}
    try:
        document = Document(io.BytesIO(data))
        unit = 0
        for paragraph in document.paragraphs:
            text = paragraph.text or ""
            if not clean(text):
                continue
            unit += 1
            units.append(
                {
                    "unit_number": unit,
                    "unit_type": "DOCX_PARAGRAPH",
                    "text": text,
                    "text_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                    "error": None,
                }
            )
        for table_number, table in enumerate(document.tables, start=1):
            for row_number, row in enumerate(table.rows, start=1):
                text = " | ".join(cell.text for cell in row.cells)
                unit += 1
                units.append(
                    {
                        "unit_number": unit,
                        "unit_type": f"DOCX_TABLE_{table_number}_ROW_{row_number}",
                        "text": text,
                        "text_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                        "error": None,
                    }
                )
        result["paragraphs"] = len(document.paragraphs)
        result["tables"] = len(document.tables)
    except Exception as exc:  # noqa: BLE001
        result["parse_error"] = f"{type(exc).__name__}: {exc}"
    return units, result


def extract_xlsx(data: bytes) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    from openpyxl import load_workbook

    units: list[dict[str, Any]] = []
    result: dict[str, Any] = {"sheets": []}
    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=False)
        unit = 0
        for worksheet in workbook.worksheets:
            result["sheets"].append(
                {"name": worksheet.title, "max_row": worksheet.max_row, "max_column": worksheet.max_column}
            )
            for row_number, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value) for value in row]
                if not any(clean(value) for value in values):
                    continue
                unit += 1
                text = "\t".join(values)
                units.append(
                    {
                        "unit_number": unit,
                        "unit_type": f"XLSX:{worksheet.title}:ROW:{row_number}",
                        "text": text,
                        "text_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                        "error": None,
                    }
                )
    except Exception as exc:  # noqa: BLE001
        result["parse_error"] = f"{type(exc).__name__}: {exc}"
    return units, result


def extract_zip(data: bytes) -> dict[str, Any]:
    result: dict[str, Any] = {"members": []}
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            for info in archive.infolist():
                result["members"].append(
                    {
                        "name": info.filename,
                        "uncompressed_bytes": info.file_size,
                        "compressed_bytes": info.compress_size,
                        "crc": info.CRC,
                        "is_dir": info.is_dir(),
                    }
                )
    except Exception as exc:  # noqa: BLE001
        result["parse_error"] = f"{type(exc).__name__}: {exc}"
    return result


def extract_asset(data: bytes, final_url: str, content_type: str | None) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    ext = extension_for(final_url, content_type)
    mime = (content_type or mimetypes.guess_type(final_url)[0] or "").lower()
    units: list[dict[str, Any]] = []
    structure: dict[str, Any] = {"extension": ext, "mime_type": mime}
    if ext == ".pdf" or "application/pdf" in mime:
        units, parsed = extract_pdf(data)
        structure.update(parsed)
    elif ext in {".jpg", ".jpeg", ".png", ".gif", ".webp", ".tif", ".tiff", ".bmp"} or mime.startswith("image/"):
        structure.update(extract_image(data))
    elif ext == ".docx" or "wordprocessingml" in mime:
        units, parsed = extract_docx(data)
        structure.update(parsed)
    elif ext in {".xlsx", ".xlsm"} or "spreadsheetml" in mime:
        units, parsed = extract_xlsx(data)
        structure.update(parsed)
    elif ext in {".csv", ".txt", ".rtf", ".xml", ".json", ".kml"} or mime.startswith("text/"):
        text, encoding = decode_text(data)
        units.append(
            {
                "unit_number": 1,
                "unit_type": "TEXT_FILE",
                "text": text,
                "text_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                "error": None,
            }
        )
        structure["encoding"] = encoding
    elif ext in {".zip", ".kmz"} or "zip" in mime:
        structure.update(extract_zip(data))
    return units, structure


def locate_page_db(root: Path) -> Path:
    matches = list(root.rglob("GEIPAN_CASE_PAGE_SOURCE_SNAPSHOT_v0.2.0.sqlite"))
    if len(matches) != 1:
        raise RuntimeError(f"Expected one page snapshot database, found {len(matches)}")
    return matches[0]


def plan(args: argparse.Namespace) -> None:
    root = Path(args.case_page_snapshot)
    db = locate_page_db(root)
    con = sqlite3.connect(f"file:{db}?mode=ro", uri=True)
    con.row_factory = sqlite3.Row
    rows = [
        dict(row)
        for row in con.execute(
            """
            SELECT asset_url,
                   MIN(inferred_kind) AS inferred_kind,
                   COUNT(DISTINCT case_url) AS linked_case_count,
                   GROUP_CONCAT(DISTINCT case_url) AS case_urls,
                   GROUP_CONCAT(DISTINCT link_label) AS link_labels
            FROM case_asset_link
            GROUP BY asset_url
            ORDER BY asset_url
            """
        )
    ]
    con.close()
    out = Path(args.output_dir)
    out.mkdir(parents=True, exist_ok=True)
    (out / "assets.json").write_text(json.dumps(rows, indent=2, ensure_ascii=False), encoding="utf-8")
    include = []
    for start in range(0, len(rows), args.batch_size):
        include.append(
            {
                "batch_index": start // args.batch_size,
                "start": start,
                "end": min(start + args.batch_size, len(rows)),
                "expected_assets": min(args.batch_size, len(rows) - start),
            }
        )
    matrix = {"include": include}
    (out / "matrix.json").write_text(json.dumps(matrix, separators=(",", ":")), encoding="utf-8")
    (out / "PLAN_SUMMARY.json").write_text(
        json.dumps(
            {
                "status": "PASS",
                "unique_asset_urls": len(rows),
                "batch_size": args.batch_size,
                "batch_count": len(include),
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    print(json.dumps(matrix, separators=(",", ":")))


def acquire(args: argparse.Namespace) -> None:
    assets = json.loads(Path(args.assets).read_text(encoding="utf-8"))
    subset = assets[args.start:args.end]
    out = Path(args.output_dir)
    if out.exists():
        shutil.rmtree(out)
    out.mkdir(parents=True)
    blobs = out / "blobs"
    blobs.mkdir()
    db = out / f"GEIPAN_ASSET_BATCH_{args.batch_index:04d}.sqlite"
    con = sqlite3.connect(db)
    con.execute("PRAGMA foreign_keys=ON")
    con.executescript("""
    CREATE TABLE asset(
      asset_url TEXT PRIMARY KEY,
      inferred_kind TEXT,
      linked_case_count INTEGER NOT NULL,
      case_urls TEXT,
      link_labels TEXT,
      retrieval_status TEXT NOT NULL,
      retrieved_at TEXT,
      final_url TEXT,
      http_headers_json TEXT,
      mime_type TEXT,
      byte_count INTEGER,
      sha256 TEXT,
      local_blob_path TEXT,
      structure_json TEXT,
      error TEXT
    );
    CREATE TABLE asset_text_unit(
      unit_id INTEGER PRIMARY KEY,
      asset_url TEXT NOT NULL REFERENCES asset(asset_url),
      unit_number INTEGER NOT NULL,
      unit_type TEXT NOT NULL,
      text TEXT,
      text_sha256 TEXT,
      error TEXT,
      UNIQUE(asset_url,unit_number,unit_type)
    );
    CREATE TABLE asset_issue(
      issue_id INTEGER PRIMARY KEY,
      asset_url TEXT NOT NULL,
      issue_code TEXT NOT NULL,
      severity TEXT NOT NULL,
      detail TEXT NOT NULL
    );
    CREATE INDEX idx_asset_status ON asset(retrieval_status);
    CREATE INDEX idx_asset_sha ON asset(sha256);
    """)
    statuses = Counter()
    total_bytes = 0
    for index, row in enumerate(subset, start=1):
        url = row["asset_url"]
        try:
            data, final_url, headers = shared.request_bytes(url, retries=args.retries, timeout=600)
            byte_count = len(data)
            if byte_count > args.max_asset_bytes:
                status = "DEFERRED_OVERSIZE"
                con.execute(
                    "INSERT INTO asset VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (
                        url, row.get("inferred_kind"), int(row.get("linked_case_count") or 0),
                        row.get("case_urls"), row.get("link_labels"), status,
                        time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()), final_url,
                        json.dumps(headers, ensure_ascii=False, sort_keys=True), headers.get("Content-Type"),
                        byte_count, hashlib.sha256(data).hexdigest(), None, None,
                        f"Asset exceeds configured {args.max_asset_bytes}-byte preservation cap",
                    ),
                )
                con.execute(
                    "INSERT INTO asset_issue(asset_url,issue_code,severity,detail) VALUES(?,?,?,?)",
                    (url, "ASSET_OVERSIZE", "MEDIUM", f"{byte_count} bytes"),
                )
                statuses[status] += 1
            elif total_bytes + byte_count > args.max_batch_bytes:
                status = "DEFERRED_BATCH_CAP"
                con.execute(
                    "INSERT INTO asset VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (
                        url, row.get("inferred_kind"), int(row.get("linked_case_count") or 0),
                        row.get("case_urls"), row.get("link_labels"), status,
                        time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()), final_url,
                        json.dumps(headers, ensure_ascii=False, sort_keys=True), headers.get("Content-Type"),
                        byte_count, hashlib.sha256(data).hexdigest(), None, None,
                        f"Batch exceeds configured {args.max_batch_bytes}-byte preservation cap",
                    ),
                )
                con.execute(
                    "INSERT INTO asset_issue(asset_url,issue_code,severity,detail) VALUES(?,?,?,?)",
                    (url, "BATCH_BYTE_CAP", "MEDIUM", f"{byte_count} bytes"),
                )
                statuses[status] += 1
            else:
                digest = hashlib.sha256(data).hexdigest()
                ext = extension_for(final_url, headers.get("Content-Type"))
                blob = blobs / f"{digest}{ext}"
                if not blob.exists():
                    blob.write_bytes(data)
                units, structure = extract_asset(data, final_url, headers.get("Content-Type"))
                con.execute(
                    "INSERT INTO asset VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (
                        url, row.get("inferred_kind"), int(row.get("linked_case_count") or 0),
                        row.get("case_urls"), row.get("link_labels"), "DOWNLOADED",
                        time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()), final_url,
                        json.dumps(headers, ensure_ascii=False, sort_keys=True), headers.get("Content-Type"),
                        byte_count, digest, f"blobs/{blob.name}",
                        json.dumps(structure, ensure_ascii=False, sort_keys=True), None,
                    ),
                )
                con.executemany(
                    "INSERT OR IGNORE INTO asset_text_unit(asset_url,unit_number,unit_type,text,text_sha256,error) VALUES(?,?,?,?,?,?)",
                    [
                        (
                            url, int(unit["unit_number"]), unit["unit_type"], unit.get("text"),
                            unit.get("text_sha256"), unit.get("error"),
                        )
                        for unit in units
                    ],
                )
                total_bytes += byte_count
                statuses["DOWNLOADED"] += 1
        except Exception as exc:  # noqa: BLE001
            detail = f"{type(exc).__name__}: {exc}"
            con.execute(
                "INSERT INTO asset(asset_url,inferred_kind,linked_case_count,case_urls,link_labels,retrieval_status,error) VALUES(?,?,?,?,?,?,?)",
                (
                    url, row.get("inferred_kind"), int(row.get("linked_case_count") or 0),
                    row.get("case_urls"), row.get("link_labels"), "DOWNLOAD_ERROR", detail,
                ),
            )
            con.execute(
                "INSERT INTO asset_issue(asset_url,issue_code,severity,detail) VALUES(?,?,?,?)",
                (url, "ASSET_DOWNLOAD_ERROR", "HIGH", detail),
            )
            statuses["DOWNLOAD_ERROR"] += 1
        con.commit()
        print(f"asset batch {args.batch_index}: {index}/{len(subset)} {dict(statuses)}", flush=True)
        time.sleep(args.delay_seconds)
    quick = con.execute("PRAGMA quick_check").fetchone()[0]
    fk = len(con.execute("PRAGMA foreign_key_check").fetchall())
    rows = con.execute("SELECT COUNT(*) FROM asset").fetchone()[0]
    text_units = con.execute("SELECT COUNT(*) FROM asset_text_unit").fetchone()[0]
    con.close()
    summary = {
        "batch_index": args.batch_index,
        "expected_assets": len(subset),
        "asset_rows": rows,
        "status_counts": dict(statuses),
        "downloaded_bytes": total_bytes,
        "text_units": text_units,
        "sqlite_quick_check": quick,
        "foreign_key_violations": fk,
    }
    (out / "BATCH_SUMMARY.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    checksums = []
    for path in sorted(out.rglob("*")):
        if path.is_file() and path.name != "SHA256SUMS.txt":
            checksums.append(f"{sha256_file(path)}  {path.relative_to(out).as_posix()}")
    (out / "SHA256SUMS.txt").write_text("\n".join(checksums) + "\n", encoding="utf-8")
    if quick != "ok" or fk or rows != len(subset):
        raise SystemExit(json.dumps(summary, indent=2))


def aggregate(args: argparse.Namespace) -> None:
    src = Path(args.input_dir)
    out = Path(args.output_dir)
    if out.exists():
        shutil.rmtree(out)
    out.mkdir(parents=True)
    db = out / "GEIPAN_LINKED_ASSET_METADATA_AND_TEXT_v0.1.0.sqlite"
    con = sqlite3.connect(db)
    con.execute("PRAGMA foreign_keys=ON")
    con.executescript("""
    CREATE TABLE asset(
      asset_url TEXT PRIMARY KEY,inferred_kind TEXT,linked_case_count INTEGER NOT NULL,
      case_urls TEXT,link_labels TEXT,retrieval_status TEXT NOT NULL,retrieved_at TEXT,
      final_url TEXT,http_headers_json TEXT,mime_type TEXT,byte_count INTEGER,sha256 TEXT,
      source_shard_name TEXT,source_shard_blob_path TEXT,structure_json TEXT,error TEXT
    );
    CREATE TABLE asset_text_unit(
      unit_id INTEGER PRIMARY KEY,asset_url TEXT NOT NULL REFERENCES asset(asset_url),
      unit_number INTEGER NOT NULL,unit_type TEXT NOT NULL,text TEXT,text_sha256 TEXT,error TEXT,
      UNIQUE(asset_url,unit_number,unit_type)
    );
    CREATE TABLE asset_issue(issue_id INTEGER PRIMARY KEY,asset_url TEXT NOT NULL,issue_code TEXT NOT NULL,severity TEXT NOT NULL,detail TEXT NOT NULL);
    CREATE INDEX idx_asset_status ON asset(retrieval_status);
    CREATE INDEX idx_asset_sha ON asset(sha256);
    """)
    batch_summaries = []
    for batch_db in sorted(src.rglob("GEIPAN_ASSET_BATCH_*.sqlite")):
        batch_name = batch_db.stem.replace("GEIPAN_ASSET_BATCH_", "geipan-asset-batch-")
        summary_path = batch_db.parent / "BATCH_SUMMARY.json"
        if summary_path.exists():
            batch_summaries.append(json.loads(summary_path.read_text(encoding="utf-8")))
        bcon = sqlite3.connect(f"file:{batch_db}?mode=ro", uri=True)
        bcon.row_factory = sqlite3.Row
        for row in bcon.execute("SELECT * FROM asset"):
            data = dict(row)
            con.execute(
                "INSERT OR REPLACE INTO asset VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (
                    data["asset_url"], data["inferred_kind"], data["linked_case_count"],
                    data["case_urls"], data["link_labels"], data["retrieval_status"],
                    data["retrieved_at"], data["final_url"], data["http_headers_json"],
                    data["mime_type"], data["byte_count"], data["sha256"],
                    batch_name, data["local_blob_path"], data["structure_json"], data["error"],
                ),
            )
        for row in bcon.execute("SELECT asset_url,unit_number,unit_type,text,text_sha256,error FROM asset_text_unit"):
            con.execute(
                "INSERT OR IGNORE INTO asset_text_unit(asset_url,unit_number,unit_type,text,text_sha256,error) VALUES(?,?,?,?,?,?)",
                tuple(row),
            )
        for row in bcon.execute("SELECT asset_url,issue_code,severity,detail FROM asset_issue"):
            con.execute("INSERT INTO asset_issue(asset_url,issue_code,severity,detail) VALUES(?,?,?,?)", tuple(row))
        bcon.close()
        con.commit()
    quick = con.execute("PRAGMA quick_check").fetchone()[0]
    fk = len(con.execute("PRAGMA foreign_key_check").fetchall())
    asset_count = con.execute("SELECT COUNT(*) FROM asset").fetchone()[0]
    status_counts = dict(con.execute("SELECT retrieval_status,COUNT(*) FROM asset GROUP BY retrieval_status"))
    downloaded_bytes = con.execute("SELECT COALESCE(SUM(byte_count),0) FROM asset WHERE retrieval_status='DOWNLOADED'").fetchone()[0]
    text_units = con.execute("SELECT COUNT(*) FROM asset_text_unit").fetchone()[0]
    issue_count = con.execute("SELECT COUNT(*) FROM asset_issue").fetchone()[0]
    con.close()
    summary = {
        "overall_status": "PASS_SOURCE_ACCOUNTED" if quick == "ok" and fk == 0 and asset_count == args.expected_assets else "FAIL",
        "expected_assets": args.expected_assets,
        "asset_rows": asset_count,
        "status_counts": status_counts,
        "downloaded_bytes": downloaded_bytes,
        "text_units": text_units,
        "issue_rows": issue_count,
        "batch_count": len(batch_summaries),
        "sqlite_quick_check": quick,
        "foreign_key_violations": fk,
        "raw_binary_policy": "Raw assets remain in named per-batch source-shard artifacts; aggregate database stores their hashes, shard paths, metadata, and extracted content.",
    }
    (out / "SUMMARY.json").write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")
    with (out / "ASSET_INDEX.csv").open("w", encoding="utf-8", newline="") as handle:
        con2 = sqlite3.connect(db)
        cur = con2.execute("SELECT asset_url,inferred_kind,retrieval_status,final_url,mime_type,byte_count,sha256,source_shard_name,source_shard_blob_path,error FROM asset ORDER BY asset_url")
        writer = csv.writer(handle)
        writer.writerow([column[0] for column in cur.description])
        writer.writerows(cur)
        con2.close()
    checksums = []
    for path in sorted(out.rglob("*")):
        if path.is_file() and path.name != "SHA256SUMS.txt":
            checksums.append(f"{sha256_file(path)}  {path.relative_to(out).as_posix()}")
    (out / "SHA256SUMS.txt").write_text("\n".join(checksums) + "\n", encoding="utf-8")
    if summary["overall_status"] == "FAIL":
        raise SystemExit(json.dumps(summary, indent=2))


def parser() -> argparse.ArgumentParser:
    root = argparse.ArgumentParser()
    commands = root.add_subparsers(dest="command", required=True)
    p = commands.add_parser("plan")
    p.add_argument("--case-page-snapshot", required=True)
    p.add_argument("--output-dir", required=True)
    p.add_argument("--batch-size", type=int, default=50)
    p.set_defaults(function=plan)
    a = commands.add_parser("acquire")
    a.add_argument("--assets", required=True)
    a.add_argument("--batch-index", type=int, required=True)
    a.add_argument("--start", type=int, required=True)
    a.add_argument("--end", type=int, required=True)
    a.add_argument("--output-dir", required=True)
    a.add_argument("--delay-seconds", type=float, default=4.0)
    a.add_argument("--retries", type=int, default=9)
    a.add_argument("--max-asset-bytes", type=int, default=500_000_000)
    a.add_argument("--max-batch-bytes", type=int, default=1_500_000_000)
    a.set_defaults(function=acquire)
    g = commands.add_parser("aggregate")
    g.add_argument("--input-dir", required=True)
    g.add_argument("--output-dir", required=True)
    g.add_argument("--expected-assets", type=int, required=True)
    g.set_defaults(function=aggregate)
    return root


def main() -> None:
    args = parser().parse_args()
    args.function(args)


if __name__ == "__main__":
    main()
